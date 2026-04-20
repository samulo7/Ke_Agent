from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
import os
import sqlite3
import uuid
from collections.abc import Sequence
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal

from app.core.env_loader import load_project_env
from app.rag.knowledge_retriever import KnowledgeRetriever
from app.rag.sample_corpus import SAMPLE_KNOWLEDGE_VERSION, load_sample_entries
from app.repos.postgres_knowledge_repository import PostgresKnowledgeRepository
from app.repos.postgres_schema import bootstrap_postgres_schema
from app.repos.runtime_connection import DBBackend, RuntimeConnection, detect_backend, wrap_runtime_connection
from app.repos.sql_knowledge_repository import SQLKnowledgeRepository, bootstrap_sqlite_schema
from app.schemas.dingtalk_chat import IntentType
from app.schemas.knowledge import KnowledgeAccessContext, KnowledgeEntry
from app.services.intent_classifier import IntentClassifier
from app.services.knowledge_answering import KnowledgeAnswerService

RoleCode = Literal["hr", "finance", "business", "admin"]
KnowledgeKind = Literal["policy_doc", "faq", "fixed_quote", "restricted_doc"]


@dataclass(frozen=True)
class QuoteFieldsInput:
    quote_item_name: str
    spec_model: str
    quote_category: str
    price_amount: float
    unit: str
    tax_included: bool
    effective_date: str
    quote_version: str
    non_standard_action: str
    quote_item_code: str = ""
    price_currency: str = "CNY"
    expire_date: str = ""
    source_note: str = ""
    has_price_conflict: bool = False
    price_conflict_note: str = ""


@dataclass(frozen=True)
class KnowledgeInput:
    knowledge_kind: KnowledgeKind
    title: str
    summary: str
    applicability: str
    next_step: str
    source_uri: str
    updated_at: str
    owner: str
    department: str
    permission_scope: str
    permitted_depts: tuple[str, ...]
    keywords: tuple[str, ...]
    intents: tuple[str, ...]
    version_tag: str
    category: str
    quote_fields: QuoteFieldsInput | None = None


@dataclass(frozen=True)
class ParsedKnowledgeDocument:
    summary: str
    source_uri: str
    chunk_texts: tuple[str, ...]
    extracted_keywords: tuple[str, ...]


_MENU_PERMISSIONS: dict[RoleCode, dict[str, bool]] = {
    "hr": {
        "dashboard": True,
        "todos": True,
        "knowledge": True,
        "import": True,
        "review": True,
        "validation": True,
        "publish": True,
        "roles": False,
    },
    "business": {
        "dashboard": True,
        "todos": True,
        "knowledge": True,
        "import": True,
        "review": True,
        "validation": True,
        "publish": True,
        "roles": False,
    },
    "finance": {
        "dashboard": True,
        "todos": False,
        "knowledge": True,
        "import": False,
        "review": False,
        "validation": True,
        "publish": False,
        "roles": False,
    },
    "admin": {
        "dashboard": True,
        "todos": True,
        "knowledge": True,
        "import": True,
        "review": True,
        "validation": True,
        "publish": True,
        "roles": True,
    },
}

_KIND_PERMISSIONS: dict[RoleCode, dict[KnowledgeKind, dict[str, bool]]] = {
    "hr": {
        "policy_doc": {"can_view": True, "can_create": True, "can_edit": True, "can_publish": True, "can_disable": True},
        "faq": {"can_view": True, "can_create": True, "can_edit": True, "can_publish": True, "can_disable": True},
        "fixed_quote": {"can_view": True, "can_create": False, "can_edit": False, "can_publish": False, "can_disable": False},
        "restricted_doc": {"can_view": True, "can_create": True, "can_edit": True, "can_publish": True, "can_disable": True},
    },
    "business": {
        "policy_doc": {"can_view": True, "can_create": False, "can_edit": False, "can_publish": False, "can_disable": False},
        "faq": {"can_view": True, "can_create": False, "can_edit": False, "can_publish": False, "can_disable": False},
        "fixed_quote": {"can_view": True, "can_create": True, "can_edit": True, "can_publish": True, "can_disable": True},
        "restricted_doc": {"can_view": True, "can_create": False, "can_edit": False, "can_publish": False, "can_disable": False},
    },
    "finance": {
        "policy_doc": {"can_view": True, "can_create": False, "can_edit": False, "can_publish": False, "can_disable": False},
        "faq": {"can_view": True, "can_create": False, "can_edit": False, "can_publish": False, "can_disable": False},
        "fixed_quote": {"can_view": True, "can_create": False, "can_edit": False, "can_publish": False, "can_disable": False},
        "restricted_doc": {"can_view": True, "can_create": False, "can_edit": False, "can_publish": False, "can_disable": False},
    },
    "admin": {
        kind: {"can_view": True, "can_create": True, "can_edit": True, "can_publish": True, "can_disable": True}
        for kind in ("policy_doc", "faq", "fixed_quote", "restricted_doc")
    },
}


class AdminKnowledgeForbiddenError(PermissionError):
    pass


class AdminKnowledgeNotFoundError(LookupError):
    pass


class AdminKnowledgeValidationError(ValueError):
    def __init__(self, message: str, *, field: str = "") -> None:
        super().__init__(message)
        self.field = field


class AdminKnowledgeService:
    def __init__(self, *, connection: Any) -> None:
        self._connection = wrap_runtime_connection(connection)
        self._intent_classifier = IntentClassifier()

    @property
    def connection(self) -> RuntimeConnection:
        return self._connection

    @property
    def backend(self) -> DBBackend:
        return self._connection.backend

    def get_permissions(self, *, user_id: str, role_code: RoleCode) -> dict[str, Any]:
        self._validate_role(role_code)
        return {
            "user_id": user_id,
            "role_code": role_code,
            "menus": dict(_MENU_PERMISSIONS[role_code]),
            "knowledge_permissions": {
                kind: dict(perms)
                for kind, perms in _KIND_PERMISSIONS[role_code].items()
            },
        }

    def list_knowledge(
        self,
        *,
        role_code: RoleCode,
        knowledge_kind: str = "",
        review_status: str = "",
        keyword: str = "",
        page: int = 1,
        page_size: int = 20,
    ) -> dict[str, Any]:
        self._validate_role(role_code)
        conditions = ["is_deleted = 0"]
        params: list[Any] = []
        if knowledge_kind:
            conditions.append("knowledge_kind = ?")
            params.append(knowledge_kind)
        if review_status:
            conditions.append("review_status = ?")
            params.append(review_status)
        if keyword:
            conditions.append("(title LIKE ? OR keywords_csv LIKE ?)")
            like = f"%{keyword}%"
            params.extend((like, like))
        where_sql = " AND ".join(conditions)
        offset = max(page - 1, 0) * page_size
        count_row = self._connection.execute(
            f"SELECT COUNT(*) AS total FROM knowledge_docs WHERE {where_sql}",
            params,
        ).fetchone()
        rows = self._connection.execute(
            f"""
            SELECT doc_id, title, knowledge_kind, source_type, review_status, owner, updated_at, published_at, last_validated_at
            FROM knowledge_docs
            WHERE {where_sql}
            ORDER BY updated_at DESC, doc_id DESC
            LIMIT ? OFFSET ?
            """,
            (*params, page_size, offset),
        ).fetchall()
        items = [self._row_to_list_item(row=row, role_code=role_code) for row in rows]
        return {
            "items": items,
            "pagination": {
                "page": page,
                "page_size": page_size,
                "total": int(count_row["total"] if count_row is not None else 0),
            },
        }

    def get_knowledge_detail(self, *, role_code: RoleCode, doc_id: str) -> dict[str, Any]:
        self._validate_role(role_code)
        row = self._get_doc(doc_id)
        if row is None:
            raise AdminKnowledgeNotFoundError(f"knowledge not found: {doc_id}")
        knowledge_kind = str(row["knowledge_kind"] or "policy_doc")
        permissions = _KIND_PERMISSIONS[role_code][knowledge_kind]
        quote_fields = None
        if knowledge_kind == "fixed_quote":
            quote_row = self._connection.execute(
                "SELECT * FROM knowledge_quote_fields WHERE doc_id = ?",
                (doc_id,),
            ).fetchone()
            if quote_row is not None:
                quote_fields = {
                    "quote_item_name": str(quote_row["quote_item_name"] or ""),
                    "quote_item_code": str(quote_row["quote_item_code"] or ""),
                    "spec_model": str(quote_row["spec_model"] or ""),
                    "quote_category": str(quote_row["quote_category"] or ""),
                    "price_amount": float(quote_row["price_amount"] or 0),
                    "price_currency": str(quote_row["price_currency"] or "CNY"),
                    "unit": str(quote_row["unit"] or ""),
                    "tax_included": bool(int(quote_row["tax_included"] or 0)),
                    "effective_date": str(quote_row["effective_date"] or ""),
                    "expire_date": str(quote_row["expire_date"] or ""),
                    "quote_version": str(quote_row["quote_version"] or ""),
                    "non_standard_action": str(quote_row["non_standard_action"] or ""),
                    "source_note": str(quote_row["source_note"] or ""),
                    "has_price_conflict": bool(int(quote_row["has_price_conflict"] or 0)),
                    "price_conflict_note": str(quote_row["price_conflict_note"] or ""),
                }
        return {
            "knowledge": {
                "doc_id": str(row["doc_id"]),
                "title": str(row["title"] or ""),
                "knowledge_kind": knowledge_kind,
                "source_type": str(row["source_type"] or "document"),
                "summary": str(row["summary"] or ""),
                "applicability": str(row["applicability"] or ""),
                "next_step": str(row["next_step"] or ""),
                "source_uri": str(row["source_uri"] or ""),
                "updated_at": str(row["updated_at"] or ""),
                "version_tag": str(row["version_tag"] or ""),
                "owner": str(row["owner"] or ""),
                "department": "",
                "review_status": str(row["review_status"] or "draft"),
                "permission_scope": str(row["permission_scope"] or "public"),
                "permitted_depts": _parse_csv(str(row["permitted_depts_csv"] or "")),
                "keywords": _parse_csv(str(row["keywords_csv"] or "")),
                "intents": _parse_csv(str(row["intents_csv"] or "")),
                "last_validated_at": str(row["last_validated_at"] or ""),
                "published_at": str(row["published_at"] or ""),
            },
            "quote_fields": quote_fields,
            "permissions": dict(permissions),
        }

    def create_knowledge(
        self,
        *,
        user_id: str,
        role_code: RoleCode,
        payload: KnowledgeInput,
    ) -> dict[str, Any]:
        self._assert_action_allowed(role_code=role_code, knowledge_kind=payload.knowledge_kind, action="can_create")
        self._validate_payload(payload)
        doc_id = f"doc-{uuid.uuid4().hex[:12]}"
        source_type = "faq" if payload.knowledge_kind in {"faq", "fixed_quote"} else "document"
        self._connection.execute(
            """
            INSERT INTO knowledge_docs (
                doc_id, source_type, title, summary, applicability, next_step, source_uri, updated_at,
                status, owner, category, version_tag, keywords_csv, intents_csv, permission_scope,
                permitted_depts_csv, knowledge_kind, review_status, created_by, updated_by,
                published_by, published_at, last_validated_at, is_deleted
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                doc_id,
                source_type,
                payload.title,
                payload.summary,
                payload.applicability,
                payload.next_step,
                payload.source_uri,
                payload.updated_at,
                "active",
                payload.owner,
                payload.category,
                payload.version_tag,
                _csv(payload.keywords),
                _csv(payload.intents),
                payload.permission_scope,
                _csv(payload.permitted_depts),
                payload.knowledge_kind,
                "draft",
                user_id,
                user_id,
                "",
                "",
                "",
                0,
            ),
        )
        self._upsert_default_chunk(doc_id=doc_id, payload=payload)
        if payload.knowledge_kind == "fixed_quote":
            assert payload.quote_fields is not None
            self._upsert_quote_fields(doc_id=doc_id, quote_fields=payload.quote_fields)
        self._connection.commit()
        return {"doc_id": doc_id, "review_status": "draft"}

    def upload_knowledge_document(
        self,
        *,
        user_id: str,
        role_code: RoleCode,
        payload: KnowledgeInput,
        filename: str,
        content: bytes,
    ) -> dict[str, Any]:
        self._assert_action_allowed(role_code=role_code, knowledge_kind=payload.knowledge_kind, action="can_create")
        if payload.knowledge_kind not in {"policy_doc", "restricted_doc"}:
            raise AdminKnowledgeValidationError("document upload only supports policy_doc or restricted_doc", field="knowledge_kind")
        parsed = self._parse_uploaded_document(filename=filename, content=content)
        parsed_payload = KnowledgeInput(
            knowledge_kind=payload.knowledge_kind,
            title=payload.title,
            summary=payload.summary.strip() or parsed.summary,
            applicability=payload.applicability,
            next_step=payload.next_step,
            source_uri=_merge_source_uri(base=payload.source_uri, uploaded=parsed.source_uri),
            updated_at=payload.updated_at,
            owner=payload.owner,
            department=payload.department,
            permission_scope=payload.permission_scope,
            permitted_depts=payload.permitted_depts,
            keywords=_merge_keywords(payload.keywords, parsed.extracted_keywords),
            intents=payload.intents,
            version_tag=payload.version_tag,
            category=payload.category,
            quote_fields=None,
        )
        self._validate_payload(parsed_payload)
        doc_id = f"doc-{uuid.uuid4().hex[:12]}"
        self._connection.execute(
            """
            INSERT INTO knowledge_docs (
                doc_id, source_type, title, summary, applicability, next_step, source_uri, updated_at,
                status, owner, category, version_tag, keywords_csv, intents_csv, permission_scope,
                permitted_depts_csv, knowledge_kind, review_status, created_by, updated_by,
                published_by, published_at, last_validated_at, is_deleted
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                doc_id,
                "document",
                parsed_payload.title,
                parsed_payload.summary,
                parsed_payload.applicability,
                parsed_payload.next_step,
                parsed_payload.source_uri,
                parsed_payload.updated_at,
                "active",
                parsed_payload.owner,
                parsed_payload.category,
                parsed_payload.version_tag,
                _csv(parsed_payload.keywords),
                _csv(parsed_payload.intents),
                parsed_payload.permission_scope,
                _csv(parsed_payload.permitted_depts),
                parsed_payload.knowledge_kind,
                "draft",
                user_id,
                user_id,
                "",
                "",
                "",
                0,
            ),
        )
        self._replace_document_chunks(doc_id=doc_id, chunk_texts=parsed.chunk_texts)
        self._connection.commit()
        return {
            "doc_id": doc_id,
            "review_status": "draft",
            "source_type": "document",
            "chunk_count": len(parsed.chunk_texts),
            "updated_at": parsed_payload.updated_at,
        }

    def update_knowledge(
        self,
        *,
        user_id: str,
        role_code: RoleCode,
        doc_id: str,
        payload: KnowledgeInput,
    ) -> dict[str, Any]:
        existing = self._get_doc(doc_id)
        if existing is None:
            raise AdminKnowledgeNotFoundError(f"knowledge not found: {doc_id}")
        knowledge_kind = str(existing["knowledge_kind"] or "policy_doc")
        self._assert_action_allowed(role_code=role_code, knowledge_kind=knowledge_kind, action="can_edit")
        if payload.knowledge_kind != knowledge_kind:
            raise AdminKnowledgeValidationError("knowledge_kind cannot change", field="knowledge_kind")
        self._validate_payload(payload)
        self._connection.execute(
            """
            UPDATE knowledge_docs
            SET title = ?,
                summary = ?,
                applicability = ?,
                next_step = ?,
                source_uri = ?,
                updated_at = ?,
                owner = ?,
                category = ?,
                version_tag = ?,
                keywords_csv = ?,
                intents_csv = ?,
                permission_scope = ?,
                permitted_depts_csv = ?,
                review_status = ?,
                updated_by = ?
            WHERE doc_id = ?
            """,
            (
                payload.title,
                payload.summary,
                payload.applicability,
                payload.next_step,
                payload.source_uri,
                payload.updated_at,
                payload.owner,
                payload.category,
                payload.version_tag,
                _csv(payload.keywords),
                _csv(payload.intents),
                payload.permission_scope,
                _csv(payload.permitted_depts),
                "draft",
                user_id,
                doc_id,
            ),
        )
        self._upsert_default_chunk(doc_id=doc_id, payload=payload)
        if knowledge_kind == "fixed_quote":
            assert payload.quote_fields is not None
            self._upsert_quote_fields(doc_id=doc_id, quote_fields=payload.quote_fields)
        self._connection.commit()
        return {"doc_id": doc_id, "review_status": "draft", "updated_at": payload.updated_at}

    def preview_dingtalk_reply(
        self,
        *,
        user_id: str,
        role_code: RoleCode,
        question: str,
        doc_id: str = "",
        dept_context: str = "",
    ) -> dict[str, Any]:
        self._validate_role(role_code)
        question = question.strip()
        if not question:
            raise AdminKnowledgeValidationError("question is required", field="question")
        target = None
        if doc_id:
            target = self._get_doc(doc_id)
            if target is None:
                raise AdminKnowledgeNotFoundError(f"knowledge not found: {doc_id}")
        repository = build_knowledge_repository(connection=self._connection)
        retriever = KnowledgeRetriever(repository=repository)
        answer_service = KnowledgeAnswerService(retriever=retriever, repository=repository)
        intent = self._intent_classifier.classify(question).intent
        candidate_intents = self._validation_candidate_intents(target=target, classified_intent=intent)
        answer = None
        answer_intent = intent
        for candidate_intent in candidate_intents:
            candidate_answer = answer_service.answer(
                question=question,
                intent=candidate_intent,
                access_context=KnowledgeAccessContext(user_id=user_id, dept_id=dept_context),
                conversation_id="admin-preview",
                sender_id=user_id,
            )
            if candidate_answer.found or candidate_answer.permission_decision in {"summary_only", "deny"}:
                answer = candidate_answer
                answer_intent = candidate_intent
                break
            if answer is None:
                answer = candidate_answer
                answer_intent = candidate_intent
        assert answer is not None
        validation_id = f"validation-{uuid.uuid4().hex[:12]}"
        validation_result = "passed" if answer.found or answer.permission_decision in {"summary_only", "deny"} else "failed"
        matched_doc_ids = list(answer.source_ids)
        top_next_step = ""
        if matched_doc_ids:
            row = self._get_doc(matched_doc_ids[0])
            top_next_step = str(row["next_step"] or "") if row is not None else ""
        reply_preview = {
            "channel": "text",
            "text": answer.text,
            "interactive_card": None,
        }
        validation_doc_id = doc_id if doc_id else (None if self.backend == "postgres" else "")
        self._connection.execute(
            """
            INSERT INTO knowledge_validation_runs (
                validation_id, doc_id, question, role_context, dept_context, matched_doc_ids_json,
                reply_channel, reply_preview_json, permission_decision, validation_result,
                validated_by, validated_at, note
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                validation_id,
                validation_doc_id,
                question,
                role_code,
                dept_context,
                _json_array(matched_doc_ids),
                reply_preview["channel"],
                _json_object(reply_preview),
                answer.permission_decision,
                validation_result,
                user_id,
                answer.answered_at,
                f"intent={answer_intent}",
            ),
        )
        if doc_id and validation_result == "passed":
            self._connection.execute(
                "UPDATE knowledge_docs SET last_validated_at = ? WHERE doc_id = ?",
                (answer.answered_at, doc_id),
            )
        self._connection.commit()
        return {
            "matched_knowledge": [
                {"doc_id": source_id, "title": citation.title, "rank": idx + 1}
                for idx, (source_id, citation) in enumerate(zip(answer.source_ids, answer.citations, strict=True))
            ],
            "reply_preview": reply_preview,
            "permission_decision": answer.permission_decision,
            "citations": [item.to_dict() for item in answer.citations],
            "next_step": top_next_step,
            "validation_result": validation_result,
            "validation_id": validation_id,
        }

    def publish_knowledge(
        self,
        *,
        user_id: str,
        role_code: RoleCode,
        doc_id: str,
        publish_note: str = "",
    ) -> dict[str, Any]:
        existing = self._get_doc(doc_id)
        if existing is None:
            raise AdminKnowledgeNotFoundError(f"knowledge not found: {doc_id}")
        knowledge_kind = str(existing["knowledge_kind"] or "policy_doc")
        self._assert_action_allowed(role_code=role_code, knowledge_kind=knowledge_kind, action="can_publish")
        self._precheck_publish(existing)
        published_at = str(existing["updated_at"] or "")
        if not published_at:
            raise AdminKnowledgeValidationError("updated_at is required before publish", field="updated_at")
        self._connection.execute(
            """
            UPDATE knowledge_docs
            SET review_status = ?,
                published_by = ?,
                published_at = ?
            WHERE doc_id = ?
            """,
            ("published", user_id, published_at, doc_id),
        )
        publish_log_id = f"publish-{uuid.uuid4().hex[:12]}"
        validation_ref = None if self.backend == "postgres" else ""
        self._connection.execute(
            """
            INSERT INTO knowledge_publish_logs (
                publish_log_id, doc_id, publish_action, publish_status, validation_id, published_by, published_at, note
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (publish_log_id, doc_id, "publish", "success", validation_ref, user_id, published_at, publish_note),
        )
        self._connection.commit()
        return {
            "doc_id": doc_id,
            "review_status": "published",
            "published_at": published_at,
            "published_by": user_id,
        }

    @staticmethod
    def _validation_candidate_intents(*, target: Any | None, classified_intent: IntentType) -> tuple[IntentType, ...]:
        candidates: list[IntentType] = [classified_intent]
        if target is not None:
            for raw_intent in _parse_csv(str(target["intents_csv"] or "")):
                if raw_intent not in candidates:
                    candidates.append(raw_intent)
        return tuple(candidates)

    def _precheck_publish(self, row: Any) -> None:
        if not str(row["last_validated_at"] or "").strip():
            raise AdminKnowledgeValidationError("knowledge has not been validated", field="last_validated_at")
        knowledge_kind = str(row["knowledge_kind"] or "policy_doc")
        if knowledge_kind == "fixed_quote":
            quote_row = self._connection.execute(
                "SELECT price_amount, unit, effective_date FROM knowledge_quote_fields WHERE doc_id = ?",
                (str(row["doc_id"]),),
            ).fetchone()
            if quote_row is None:
                raise AdminKnowledgeValidationError("quote_fields is required", field="quote_fields")
            if float(quote_row["price_amount"] or 0) <= 0:
                raise AdminKnowledgeValidationError("price_amount is required", field="price_amount")
            if not str(quote_row["unit"] or "").strip():
                raise AdminKnowledgeValidationError("unit is required", field="unit")
            if not str(quote_row["effective_date"] or "").strip():
                raise AdminKnowledgeValidationError("effective_date is required", field="effective_date")

    def _upsert_quote_fields(self, *, doc_id: str, quote_fields: QuoteFieldsInput) -> None:
        self._connection.execute(
            """
            INSERT INTO knowledge_quote_fields (
                doc_id, quote_item_name, quote_item_code, spec_model, quote_category, price_amount,
                price_currency, unit, tax_included, effective_date, expire_date, quote_version,
                non_standard_action, source_note, has_price_conflict, price_conflict_note
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(doc_id) DO UPDATE SET
                quote_item_name = excluded.quote_item_name,
                quote_item_code = excluded.quote_item_code,
                spec_model = excluded.spec_model,
                quote_category = excluded.quote_category,
                price_amount = excluded.price_amount,
                price_currency = excluded.price_currency,
                unit = excluded.unit,
                tax_included = excluded.tax_included,
                effective_date = excluded.effective_date,
                expire_date = excluded.expire_date,
                quote_version = excluded.quote_version,
                non_standard_action = excluded.non_standard_action,
                source_note = excluded.source_note,
                has_price_conflict = excluded.has_price_conflict,
                price_conflict_note = excluded.price_conflict_note
            """,
            (
                doc_id,
                quote_fields.quote_item_name,
                quote_fields.quote_item_code,
                quote_fields.spec_model,
                quote_fields.quote_category,
                quote_fields.price_amount,
                quote_fields.price_currency,
                quote_fields.unit,
                1 if quote_fields.tax_included else 0,
                quote_fields.effective_date,
                quote_fields.expire_date,
                quote_fields.quote_version,
                quote_fields.non_standard_action,
                quote_fields.source_note,
                1 if quote_fields.has_price_conflict else 0,
                quote_fields.price_conflict_note,
            ),
        )

    def _upsert_default_chunk(self, *, doc_id: str, payload: KnowledgeInput) -> None:
        chunk_text = self._chunk_text_for_payload(payload)
        self._connection.execute(
            """
            INSERT INTO doc_chunks (chunk_id, doc_id, chunk_index, chunk_text, chunk_vector)
            VALUES (?, ?, ?, ?, ?)
            ON CONFLICT(chunk_id) DO UPDATE SET
                chunk_text = excluded.chunk_text,
                chunk_vector = excluded.chunk_vector
            """,
            (f"chunk-{doc_id}-0", doc_id, 0, chunk_text, "[]"),
        )

    def _replace_document_chunks(self, *, doc_id: str, chunk_texts: Sequence[str]) -> None:
        self._connection.execute("DELETE FROM doc_chunks WHERE doc_id = ?", (doc_id,))
        for index, chunk_text in enumerate(chunk_texts):
            self._connection.execute(
                """
                INSERT INTO doc_chunks (chunk_id, doc_id, chunk_index, chunk_text, chunk_vector)
                VALUES (?, ?, ?, ?, ?)
                """,
                (f"chunk-{doc_id}-{index}", doc_id, index, chunk_text, "[]"),
            )

    def _parse_uploaded_document(self, *, filename: str, content: bytes) -> ParsedKnowledgeDocument:
        normalized_name = filename.strip()
        if not normalized_name:
            raise AdminKnowledgeValidationError("filename is required", field="file")
        suffix = Path(normalized_name).suffix.lower()
        if suffix in {".txt", ".md"}:
            text = content.decode("utf-8-sig", errors="ignore")
        elif suffix == ".docx":
            text = self._extract_docx_text(content)
        elif suffix == ".pdf":
            text = self._extract_pdf_text(content)
        else:
            raise AdminKnowledgeValidationError("unsupported file type", field="file")
        normalized_text = _normalize_document_text(text)
        if not normalized_text:
            raise AdminKnowledgeValidationError("document content is empty", field="file")
        chunk_texts = _split_document_chunks(normalized_text)
        summary = _build_document_summary(_pick_summary_source(chunk_texts))
        return ParsedKnowledgeDocument(
            summary=summary,
            source_uri=f"upload:{normalized_name}",
            chunk_texts=chunk_texts,
            extracted_keywords=_extract_document_keywords(title=Path(normalized_name).stem, chunk_texts=chunk_texts),
        )

    @staticmethod
    def _extract_docx_text(content: bytes) -> str:
        from docx import Document

        document = Document(BytesIO(content))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    @staticmethod
    def _extract_pdf_text(content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise AdminKnowledgeValidationError("pdf upload requires an installed PDF parser dependency", field="file") from exc

        reader = PdfReader(BytesIO(content))
        texts = [(page.extract_text() or "").strip() for page in reader.pages]
        extracted = "\n".join(text for text in texts if text)
        if _looks_like_scanned_pdf(extracted):
            raise AdminKnowledgeValidationError("pdf appears to be scanned or text extraction failed; OCR is not supported yet", field="file")
        return extracted

    @staticmethod
    def _chunk_text_for_payload(payload: KnowledgeInput) -> str:
        parts = [payload.title, payload.summary, payload.applicability, payload.next_step, _csv(payload.keywords)]
        if payload.quote_fields is not None:
            parts.extend(
                [
                    payload.quote_fields.quote_item_name,
                    payload.quote_fields.spec_model,
                    str(payload.quote_fields.price_amount),
                    payload.quote_fields.unit,
                    payload.quote_fields.non_standard_action,
                ]
            )
        return "\n".join(part.strip() for part in parts if part and part.strip())

    def _row_to_list_item(self, *, row: Any, role_code: RoleCode) -> dict[str, Any]:
        knowledge_kind = str(row["knowledge_kind"] or "policy_doc")
        permissions = _KIND_PERMISSIONS[role_code][knowledge_kind]
        return {
            "doc_id": str(row["doc_id"]),
            "title": str(row["title"]),
            "knowledge_kind": knowledge_kind,
            "source_type": str(row["source_type"]),
            "review_status": str(row["review_status"]),
            "owner": str(row["owner"] or ""),
            "department": "",
            "updated_at": str(row["updated_at"] or ""),
            "published_at": str(row["published_at"] or ""),
            "last_validated_at": str(row["last_validated_at"] or ""),
            "hit_trend": "-",
            **permissions,
        }

    def _get_doc(self, doc_id: str) -> Any | None:
        return self._connection.execute(
            "SELECT * FROM knowledge_docs WHERE doc_id = ? AND is_deleted = 0",
            (doc_id,),
        ).fetchone()

    @staticmethod
    def _validate_role(role_code: str) -> None:
        if role_code not in _MENU_PERMISSIONS:
            raise AdminKnowledgeValidationError("invalid role_code", field="role_code")

    def _assert_action_allowed(self, *, role_code: RoleCode, knowledge_kind: KnowledgeKind, action: str) -> None:
        self._validate_role(role_code)
        permissions = _KIND_PERMISSIONS[role_code][knowledge_kind]
        if not permissions[action]:
            raise AdminKnowledgeForbiddenError(f"current role cannot {action} {knowledge_kind}")

    @staticmethod
    def _validate_payload(payload: KnowledgeInput) -> None:
        if not payload.title.strip():
            raise AdminKnowledgeValidationError("title is required", field="title")
        if not payload.summary.strip():
            raise AdminKnowledgeValidationError("summary is required", field="summary")
        if not payload.owner.strip():
            raise AdminKnowledgeValidationError("owner is required", field="owner")
        if not payload.updated_at.strip():
            raise AdminKnowledgeValidationError("updated_at is required", field="updated_at")
        if payload.knowledge_kind == "faq":
            if not payload.next_step.strip():
                raise AdminKnowledgeValidationError("next_step is required", field="next_step")
            return
        if payload.knowledge_kind == "fixed_quote":
            if payload.quote_fields is None:
                raise AdminKnowledgeValidationError("quote_fields is required", field="quote_fields")
            if payload.quote_fields.price_amount <= 0:
                raise AdminKnowledgeValidationError("price_amount is required", field="price_amount")
            if not payload.quote_fields.unit.strip():
                raise AdminKnowledgeValidationError("unit is required", field="unit")
            if not payload.quote_fields.effective_date.strip():
                raise AdminKnowledgeValidationError("effective_date is required", field="effective_date")
            if not payload.quote_fields.non_standard_action.strip():
                raise AdminKnowledgeValidationError("non_standard_action is required", field="non_standard_action")


def build_shared_admin_runtime_services() -> tuple[AdminKnowledgeService, KnowledgeAnswerService]:
    connection = connect_runtime_db()
    _seed_sample_entries_if_empty(connection)
    admin_service = AdminKnowledgeService(connection=connection)
    repository = build_knowledge_repository(connection=connection, version=SAMPLE_KNOWLEDGE_VERSION)
    answer_service = KnowledgeAnswerService(
        retriever=KnowledgeRetriever(repository=repository),
        repository=repository,
    )
    return admin_service, answer_service


def build_default_admin_knowledge_service() -> AdminKnowledgeService:
    admin_service, _ = build_shared_admin_runtime_services()
    return admin_service


def build_knowledge_repository(*, connection: Any, version: str = "a10-sql-v1") -> SQLKnowledgeRepository | PostgresKnowledgeRepository:
    runtime = wrap_runtime_connection(connection)
    if runtime.backend == "postgres":
        return PostgresKnowledgeRepository(connection=runtime, version=version)
    return SQLKnowledgeRepository(connection=runtime.raw, version=version)


def connect_shared_runtime_db() -> RuntimeConnection:
    return connect_runtime_db()


def connect_runtime_db() -> RuntimeConnection:
    explicit_backend = os.getenv("KEAGENT_DB_BACKEND")
    explicit_sqlite_path = os.getenv("KEAGENT_SQLITE_PATH")
    load_project_env()
    backend_value = explicit_backend if explicit_backend is not None else ("sqlite" if explicit_sqlite_path else os.getenv("KEAGENT_DB_BACKEND"))
    backend = _resolve_db_backend(backend_value)
    if backend == "postgres":
        return connect_postgres_runtime_db()
    return connect_sqlite_runtime_db()


def connect_sqlite_runtime_db() -> RuntimeConnection:
    raw_path = (os.getenv("KEAGENT_SQLITE_PATH") or "").strip()
    if raw_path:
        db_path = Path(raw_path)
    else:
        db_path = Path(__file__).resolve().parents[2] / ".local" / "keagent_runtime.sqlite3"
    db_path.parent.mkdir(parents=True, exist_ok=True)
    connection = sqlite3.connect(str(db_path), check_same_thread=False)
    connection.row_factory = sqlite3.Row
    bootstrap_sqlite_schema(connection)
    return RuntimeConnection(backend="sqlite", raw=connection)


def connect_postgres_runtime_db() -> RuntimeConnection:
    import psycopg
    from psycopg.rows import dict_row

    host = (os.getenv("PG_HOST") or "").strip()
    database = (os.getenv("PG_DATABASE") or "").strip()
    user = (os.getenv("PG_USER") or "").strip()
    password = (os.getenv("PG_PASSWORD") or "").strip()
    port = int((os.getenv("PG_PORT") or "5432").strip() or "5432")
    if not host or not database or not user or not password:
        raise AdminKnowledgeValidationError("PG_* config is required for postgres backend", field="PG_HOST")
    connection = psycopg.connect(
        host=host,
        port=port,
        dbname=database,
        user=user,
        password=password,
        row_factory=dict_row,
    )
    bootstrap_postgres_schema(connection)
    return RuntimeConnection(backend="postgres", raw=connection)


def _seed_sample_entries_if_empty(connection: Any) -> None:
    runtime = wrap_runtime_connection(connection)
    row = runtime.execute("SELECT COUNT(*) AS total FROM knowledge_docs").fetchone()
    total = int((row or {}).get("total", 0) if isinstance(row, dict) else row["total"] if row is not None else 0)
    if total > 0:
        return
    for index, entry in enumerate(load_sample_entries()):
        knowledge_kind = _infer_knowledge_kind(entry)
        runtime.execute(
            """
            INSERT INTO knowledge_docs (
                doc_id, source_type, title, summary, applicability, next_step, source_uri, updated_at,
                status, owner, category, version_tag, keywords_csv, intents_csv, permission_scope,
                permitted_depts_csv, knowledge_kind, review_status, created_by, updated_by,
                published_by, published_at, last_validated_at, is_deleted
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                entry.source_id,
                entry.source_type,
                entry.title,
                entry.summary,
                entry.applicability,
                entry.next_step,
                entry.source_uri,
                entry.updated_at,
                "active",
                "system",
                knowledge_kind,
                SAMPLE_KNOWLEDGE_VERSION,
                _csv(entry.keywords),
                _csv(entry.intents),
                "public",
                "",
                knowledge_kind,
                "published",
                "system",
                "system",
                "system",
                entry.updated_at,
                entry.updated_at,
                0,
            ),
        )
        runtime.execute(
            """
            INSERT INTO doc_chunks (chunk_id, doc_id, chunk_index, chunk_text, chunk_vector)
            VALUES (?, ?, ?, ?, ?)
            """,
            (
                f"sample-chunk-{index}",
                entry.source_id,
                0,
                "\n".join([entry.title, entry.summary, entry.applicability, entry.next_step, _csv(entry.keywords)]),
                "[]",
            ),
        )
    runtime.commit()


def _resolve_db_backend(raw_value: str | None) -> DBBackend:
    normalized = (raw_value or "sqlite").strip().lower()
    if normalized == "postgres":
        return "postgres"
    return "sqlite"


def _infer_knowledge_kind(entry: KnowledgeEntry) -> str:
    if "fixed_quote" in entry.intents:
        return "fixed_quote"
    if entry.source_type == "faq":
        return "faq"
    return "policy_doc"


def _csv(items: Sequence[str]) -> str:
    return ",".join(item.strip() for item in items if item.strip())


def _normalize_document_text(raw: str) -> str:
    lines = [line.strip() for line in raw.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compact = [line for line in lines if line]
    return "\n".join(compact)


def _split_document_chunks(raw: str, *, max_chars: int = 800) -> tuple[str, ...]:
    paragraphs = [paragraph.strip() for paragraph in raw.split("\n") if paragraph.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        candidate = paragraph if not current else f"{current}\n{paragraph}"
        if current and len(candidate) > max_chars:
            chunks.append(current)
            current = paragraph
            continue
        current = candidate
    if current:
        chunks.append(current)
    return tuple(chunks) if chunks else (raw[:max_chars],)


def _pick_summary_source(chunk_texts: Sequence[str]) -> str:
    def score(text: str) -> int:
        normalized = text.replace("\n", " ")
        cjk_count = sum(1 for char in normalized if "\u4e00" <= char <= "\u9fff")
        alpha_num_count = sum(1 for char in normalized if char.isalnum())
        dot_leader_penalty = normalized.count(".") + normalized.count("·") + normalized.count("…")
        toc_penalty = 20 if "目录" in normalized else 0
        return cjk_count * 2 + alpha_num_count - dot_leader_penalty * 3 - toc_penalty

    return max(chunk_texts, key=score) if chunk_texts else ""


def _extract_document_keywords(*, title: str, chunk_texts: Sequence[str], max_keywords: int = 20) -> tuple[str, ...]:
    import re
    from collections import Counter

    counter: Counter[str] = Counter()
    seed_text = "\n".join((title, *chunk_texts))
    for token in re.findall(r"[\u4e00-\u9fff]{2,8}|[A-Za-z0-9][A-Za-z0-9._-]{1,31}", seed_text):
        normalized = token.strip()
        if not normalized or normalized.isdigit() or normalized in {"目录", "员工", "手册", "制度", "说明", "流程"}:
            continue
        counter[normalized] += 1
    ranked = sorted(counter.items(), key=lambda item: (-item[1], -len(item[0]), item[0]))
    return tuple(token for token, _ in ranked[:max_keywords])


def _looks_like_scanned_pdf(raw: str) -> bool:
    normalized = _normalize_document_text(raw)
    if not normalized:
        return True
    signal_count = sum(1 for char in normalized if char.isalnum() or "\u4e00" <= char <= "\u9fff")
    return signal_count < 8


def _build_document_summary(raw: str, *, max_chars: int = 220) -> str:
    summary = raw.replace("\n", " ").strip()
    if len(summary) <= max_chars:
        return summary
    return summary[: max_chars - 1].rstrip() + "…"


def _merge_keywords(primary: Sequence[str], extracted: Sequence[str], *, limit: int = 20) -> tuple[str, ...]:
    merged: list[str] = []
    for item in (*primary, *extracted):
        normalized = item.strip()
        if not normalized or normalized in merged:
            continue
        merged.append(normalized)
        if len(merged) >= limit:
            break
    return tuple(merged)


def _merge_source_uri(*, base: str, uploaded: str) -> str:
    normalized_base = base.strip()
    normalized_uploaded = uploaded.strip()
    if normalized_base and normalized_uploaded:
        return f"{normalized_base}|{normalized_uploaded}"
    return normalized_base or normalized_uploaded


def _parse_csv(raw: str) -> list[str]:
    return [item.strip() for item in raw.split(",") if item.strip()]


def _json_array(items: Sequence[str]) -> str:
    escaped = [item.replace('"', '\\"') for item in items]
    return "[" + ",".join(f'"{item}"' for item in escaped) + "]"


def _json_object(payload: dict[str, Any]) -> str:
    channel = str(payload.get("channel", "text")).replace('"', '\\"')
    text = str(payload.get("text", "")).replace('"', '\\"').replace("\n", "\\n")
    return f'{{"channel":"{channel}","text":"{text}","interactive_card":null}}'
